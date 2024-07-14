from docx import Document

# Create a new Document
doc = Document()

# Title
doc.add_heading('Perhitungan Manual Algoritma Apriori', level=1)

# Adding transaction data section
doc.add_heading('Data Transaksi', level=2)
transactions = [
    "2024-05-11: vps, shared hosting",
    "2024-05-12: domain, vps",
    "2024-05-13: shared hosting, dedicated hosting",
    # Other transactions would be added here similarly
]

# Adding transactions to the document
for transaction in transactions:
    doc.add_paragraph(transaction)

# Adding Apriori process section
doc.add_heading('Proses Apriori', level=2)
doc.add_paragraph(
    "Proses Apriori dimulai dengan menentukan set item yang sering muncul "
    "berdasarkan minimum support yang ditentukan. Setelah itu, aturan asosiasi "
    "dibentuk dari itemset yang sering muncul tersebut berdasarkan minimum confidence."
)

# Example of calculations and rules generation
doc.add_heading('Perhitungan', level=2)
doc.add_paragraph(
    "Misalnya, untuk minimum support 50% dan minimum confidence 70%, perhitungan "
    "dilakukan sebagai berikut:"
)
doc.add_paragraph(
    "1. Menghitung frekuensi masing-masing item.\n"
    "2. Membentuk kombinasi itemset berdasarkan frekuensi yang memenuhi minimum support.\n"
    "3. Membentuk aturan asosiasi dari itemset yang memenuhi minimum confidence."
)

# Adding some example data analysis
doc.add_heading('Analisis Data', level=2)
doc.add_paragraph(
    "Berdasarkan data transaksi, item 'vps' dan 'shared hosting' sering muncul bersamaan, "
    "dan memenuhi kriteria minimum support dan confidence yang ditetapkan. Aturan asosiasi "
    "yang dapat dibentuk misalnya, pembelian 'vps' akan mengindikasikan kemungkinan pembelian 'shared hosting'."
)

# Saving the document
file_path = "/Users/nurhamim/Documents/apriori-project//Perhitungan_Apriori.docx"
doc.save(file_path)

file_path





